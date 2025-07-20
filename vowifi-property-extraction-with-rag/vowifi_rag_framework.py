#!/usr/bin/env python
import os
import logging
import gc
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from pathlib import Path
from functools import lru_cache
from concurrent.futures import ThreadPoolExecutor, as_completed
import pickle
import hashlib

# Core libraries
import torch
import numpy as np
import pandas as pd
from tqdm import tqdm

# Document processing
import fitz  # PyMuPDF
import pdfplumber
from docx import Document

# LangChain components
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document as LangchainDocument
from langchain.llms import HuggingFacePipeline
from langchain.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough

# Transformers and model handling
import transformers
from transformers import (
    AutoTokenizer, 
    AutoModelForCausalLM, 
    BitsAndBytesConfig,
    pipeline
)

# Evaluation
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import difflib


@dataclass
class RAGConfig:
    """Configuration for the VoWiFi RAG Pipeline"""
    # Model configuration
    model_name: str = "HuggingFaceH4/zephyr-7b-beta"
    embedding_model: str = "BAAI/bge-base-en-v1.5"
    torch_dtype: str = "bfloat16"
    quantization: bool = True
    device_map: str = "auto"
    
    # Text processing
    chunk_size: int = 512
    chunk_overlap: int = 50
    max_new_tokens: int = 400
    temperature: float = 0.2
    
    # Retrieval settings
    retriever_k: int = 4
    similarity_threshold: float = 0.7
    
    # Property extraction
    min_property_length: int = 20
    property_patterns: List[str] = None
    
    # Performance optimization settings
    batch_size: int = 32
    max_workers: int = 4
    use_parallel_processing: bool = True
    lazy_loading: bool = True
    memory_optimization: bool = True
    cache_embeddings: bool = True
    
    # File paths
    documents_dir: str = "rfcs"
    vector_store_path: str = "vowifi_vectorstore"
    output_dir: str = "extracted_properties"
    cache_dir: str = ".cache"
    
    def __post_init__(self):
        if self.property_patterns is None:
            self.property_patterns = [
                "shall", "must", "MUST", "if", "IF",
                "includes", "sends", "contains", "provided",
                "acceptable", "unable", "required", "proposed",
                "authentication", "authorization", "security"
            ]
        
        # Optimize batch size based on available memory
        if self.memory_optimization:
            self.batch_size = min(self.batch_size, 16)  # Conservative for GPU memory
            
        # Adjust workers based on CPU count
        if self.use_parallel_processing:
            import multiprocessing
            cpu_count = multiprocessing.cpu_count()
            self.max_workers = min(self.max_workers, cpu_count)


class DocumentProcessor:
    """Advanced document processing for VoWiFi specifications"""
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self.logger = logging.getLogger(__name__)
        self._cache_dir = Path(config.cache_dir) / "documents"
        self._cache_dir.mkdir(parents=True, exist_ok=True)
    
    def _get_cache_path(self, file_path: str) -> Path:
        """Generate cache path for processed document"""
        file_hash = hashlib.md5(str(file_path).encode()).hexdigest()
        return self._cache_dir / f"{file_hash}.pkl"
    
    def _load_from_cache(self, file_path: str) -> Optional[str]:
        """Load processed text from cache if available and valid"""
        cache_path = self._get_cache_path(file_path)
        if not cache_path.exists():
            return None
            
        try:
            file_stat = os.stat(file_path)
            cache_stat = os.stat(cache_path)
            
            # Check if cache is newer than the file
            if cache_stat.st_mtime >= file_stat.st_mtime:
                with open(cache_path, 'rb') as f:
                    return pickle.load(f)
        except Exception as e:
            self.logger.warning(f"Failed to load cache for {file_path}: {e}")
        
        return None
    
    def _save_to_cache(self, file_path: str, content: str):
        """Save processed text to cache"""
        try:
            cache_path = self._get_cache_path(file_path)
            with open(cache_path, 'wb') as f:
                pickle.dump(content, f)
        except Exception as e:
            self.logger.warning(f"Failed to save cache for {file_path}: {e}")
    
    def extract_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF with multiple fallback strategies"""
        # Check cache first
        if self.config.lazy_loading:
            cached_content = self._load_from_cache(pdf_path)
            if cached_content is not None:
                return cached_content
        
        text = ""
        
        # Strategy 1: Try PyMuPDF first (best for general text)
        try:
            with fitz.open(pdf_path) as doc:
                for page in doc:
                    text += page.get_text() + "\n"
            if text.strip():
                self._save_to_cache(pdf_path, text)
                return text
        except Exception as e:
            self.logger.warning(f"PyMuPDF failed for {pdf_path}: {e}")
        
        # Strategy 2: Try pdfplumber (better for structured content)
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                self._save_to_cache(pdf_path, text)
                return text
        except Exception as e:
            self.logger.warning(f"pdfplumber failed for {pdf_path}: {e}")
        
        return text
    
    def extract_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX file"""
        # Check cache first
        if self.config.lazy_loading:
            cached_content = self._load_from_cache(docx_path)
            if cached_content is not None:
                return cached_content
        
        try:
            doc = Document(docx_path)
            content = '\n'.join([para.text for para in doc.paragraphs])
            self._save_to_cache(docx_path, content)
            return content
        except Exception as e:
            self.logger.error(f"Error extracting from DOCX {docx_path}: {e}")
            return ""
    
    def extract_from_txt(self, txt_path: str) -> str:
        """Extract text from TXT file"""
        # Check cache first
        if self.config.lazy_loading:
            cached_content = self._load_from_cache(txt_path)
            if cached_content is not None:
                return cached_content
        
        try:
            with open(txt_path, 'r', encoding='utf-8') as f:
                content = f.read()
            self._save_to_cache(txt_path, content)
            return content
        except Exception as e:
            self.logger.error(f"Error extracting from TXT {txt_path}: {e}")
            return ""
    
    def _process_single_document(self, file_path: str) -> Optional[LangchainDocument]:
        """Process a single document and return LangChain Document"""
        file_path = Path(file_path)
        
        if file_path.suffix.lower() == '.pdf':
            content = self.extract_from_pdf(str(file_path))
        elif file_path.suffix.lower() == '.docx':
            content = self.extract_from_docx(str(file_path))
        elif file_path.suffix.lower() == '.txt':
            content = self.extract_from_txt(str(file_path))
        else:
            self.logger.warning(f"Unsupported file format: {file_path}")
            return None
        
        if not content.strip():
            self.logger.warning(f"No content extracted from {file_path}")
            return None
        
        metadata = {
            "source": str(file_path),
            "filename": file_path.name,
            "file_type": file_path.suffix.lower(),
            "char_count": len(content)
        }
        
        return LangchainDocument(page_content=content, metadata=metadata)
    
    def process_document(self, file_path: str) -> LangchainDocument:
        """Process a single document and return LangChain Document (legacy method)"""
        return self._process_single_document(file_path)
    
    def process_directory(self, directory_path: str) -> List[LangchainDocument]:
        """Process all documents in a directory with parallel processing"""
        documents = []
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            self.logger.error(f"Directory does not exist: {directory_path}")
            return documents
        
        supported_extensions = {'.pdf', '.docx', '.txt'}
        files = [f for f in directory_path.rglob('*') 
                if f.is_file() and f.suffix.lower() in supported_extensions]
        
        if not files:
            self.logger.warning(f"No supported files found in {directory_path}")
            return documents
        
        # Process documents in parallel if enabled
        if self.config.use_parallel_processing and len(files) > 1:
            with ThreadPoolExecutor(max_workers=self.config.max_workers) as executor:
                future_to_file = {
                    executor.submit(self._process_single_document, str(file_path)): file_path 
                    for file_path in files
                }
                
                for future in tqdm(as_completed(future_to_file), 
                                 total=len(files), desc="Processing documents"):
                    file_path = future_to_file[future]
                    try:
                        doc = future.result()
                        if doc:
                            documents.append(doc)
                    except Exception as e:
                        self.logger.error(f"Error processing {file_path}: {e}")
        else:
            # Sequential processing for small file counts or when parallel processing is disabled
            for file_path in tqdm(files, desc="Processing documents"):
                doc = self._process_single_document(str(file_path))
                if doc:
                    documents.append(doc)
        
        self.logger.info(f"Processed {len(documents)} documents from {directory_path}")
        return documents


class VoWiFiPropertyExtractor:
    """Specialized property extraction for VoWiFi specifications"""
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self.logger = logging.getLogger(__name__)
        
        # Enhanced property patterns for VoWiFi domain
        self.property_indicators = {
            'requirements': ['shall', 'must', 'MUST', 'required', 'mandatory'],
            'conditions': ['if', 'IF', 'when', 'WHEN', 'unless', 'provided that'],
            'behaviors': ['sends', 'transmits', 'receives', 'processes', 'includes'],
            'constraints': ['unable', 'cannot', 'forbidden', 'not allowed'],
            'specifications': ['contains', 'comprises', 'consists of', 'defined as'],
            'security': ['authentication', 'authorization', 'encrypt', 'decrypt', 'certificate']
        }
        
        # Cache for property checks
        self._property_check_cache = {}
    
    @lru_cache(maxsize=1000)
    def is_potential_property(self, text: str) -> bool:
        """Check if text contains property indicators (cached for performance)"""
        text_lower = text.lower()
        
        # Check for property patterns
        for category, patterns in self.property_indicators.items():
            if any(pattern.lower() in text_lower for pattern in patterns):
                return True
        
        # Check for specific VoWiFi terms
        vowifi_terms = ['ue', 'epdg', 'ike', 'ipsec', 'esp', 'pdn', 'aaa', 'msk']
        if any(term in text_lower for term in vowifi_terms):
            return True
        
        return False
    
    def extract_properties_from_text(self, text: str) -> List[str]:
        """Extract properties from raw text with optimized processing"""
        properties = []
        
        # Use more efficient text splitting
        sentences = self._split_into_sentences(text)
        
        # Use batch processing for large texts
        if len(sentences) > self.config.batch_size:
            for i in range(0, len(sentences), self.config.batch_size):
                batch = sentences[i:i + self.config.batch_size]
                batch_properties = self._process_sentence_batch(batch)
                properties.extend(batch_properties)
        else:
            properties = self._process_sentence_batch(sentences)
        
        return properties
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Efficiently split text into sentences"""
        # Replace newlines with spaces for better sentence detection
        text = text.replace('\n', ' ')
        
        # Split on common sentence endings
        sentences = []
        current = ""
        
        for char in text:
            current += char
            if char in '.!?':
                sentence = current.strip()
                if sentence:
                    sentences.append(sentence)
                current = ""
        
        # Add remaining text if any
        if current.strip():
            sentences.append(current.strip())
            
        return sentences
    
    def _process_sentence_batch(self, sentences: List[str]) -> List[str]:
        """Process a batch of sentences for property extraction"""
        properties = []
        
        for sentence in sentences:
            # Filter by length first (fastest check)
            if len(sentence) < self.config.min_property_length:
                continue
                
            # Check if it's a potential property
            if self.is_potential_property(sentence):
                # Clean and normalize the sentence
                cleaned = self._clean_property_text(sentence)
                if cleaned:
                    properties.append(cleaned)
        
        return properties
    
    def _clean_property_text(self, text: str) -> str:
        """Clean and normalize property text with better performance"""
        # Remove extra whitespace (faster than multiple split/join)
        text = ' '.join(text.split())
        
        # Remove common artifacts
        text = text.replace('\x00', '').replace('\ufffd', '')
        
        # Ensure proper sentence ending
        if not text.endswith('.'):
            text += '.'
        
        return text if len(text) > self.config.min_property_length else ""


class VoWiFiRAGPipeline:
    """Main RAG Pipeline for VoWiFi Property Extraction"""
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self.logger = self._setup_logging()
        
        # Initialize components
        self.document_processor = DocumentProcessor(config)
        self.property_extractor = VoWiFiPropertyExtractor(config)
        
        # RAG components (initialized later)
        self.embeddings = None
        self.vectorstore = None
        self.llm = None
        self.rag_chain = None
        self.text_splitter = None
        
        # Create necessary directories
        os.makedirs(self.config.output_dir, exist_ok=True)
        os.makedirs(self.config.cache_dir, exist_ok=True)
    
    def _setup_logging(self) -> logging.Logger:
        """Setup logging configuration"""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler('vowifi_rag.log'),
                logging.StreamHandler()
            ]
        )
        return logging.getLogger(__name__)
    
    def _monitor_memory(self):
        """Monitor memory usage and perform cleanup if needed"""
        if self.config.memory_optimization:
            # Force garbage collection
            gc.collect()
            
            # Clear CUDA cache if using GPU
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
    
    def initialize_models(self):
        """Initialize embedding and language models with optimization"""
        self.logger.info("Initializing models...")
        
        try:
            # Initialize embeddings with caching
            cache_path = Path(self.config.cache_dir) / "embeddings"
            cache_path.mkdir(exist_ok=True)
            
            self.embeddings = HuggingFaceEmbeddings(
                model_name=self.config.embedding_model,
                cache_folder=str(cache_path)
            )
            
            # Initialize text splitter
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.config.chunk_size,
                chunk_overlap=self.config.chunk_overlap,
                separators=["\n\n", "\n", ".", "!", "?", " ", ""]
            )
            
            # Initialize LLM
            self._initialize_llm()
            
            self.logger.info("Models initialized successfully")
            
        except Exception as e:
            self.logger.error(f"Failed to initialize models: {e}")
            raise
        
        # Monitor memory after initialization
        self._monitor_memory()
    
    def _initialize_llm(self):
        """Initialize the language model with quantization if specified"""
        model_kwargs = {}
        
        if self.config.quantization:
            quantization_config = BitsAndBytesConfig(
                load_in_4bit=True,
                bnb_4bit_use_double_quant=True,
                bnb_4bit_quant_type="nf4",
                bnb_4bit_compute_dtype=getattr(torch, self.config.torch_dtype)
            )
            model_kwargs["quantization_config"] = quantization_config
        
        try:
            # Load model and tokenizer
            model = AutoModelForCausalLM.from_pretrained(
                self.config.model_name,
                device_map=self.config.device_map,
                **model_kwargs
            )
            
            tokenizer = AutoTokenizer.from_pretrained(self.config.model_name)
            if tokenizer.pad_token is None:
                tokenizer.pad_token = tokenizer.eos_token
            
            # Create pipeline
            text_generation_pipeline = pipeline(
                "text-generation",
                model=model,
                tokenizer=tokenizer,
                temperature=self.config.temperature,
                do_sample=True,
                repetition_penalty=1.1,
                return_full_text=False,
                max_new_tokens=self.config.max_new_tokens,
            )
            
            self.llm = HuggingFacePipeline(pipeline=text_generation_pipeline)
            
        except Exception as e:
            self.logger.error(f"Failed to initialize LLM: {e}")
            # Fallback to a lighter model or disable LLM functionality
            self.llm = None
    
    def build_knowledge_base(self, documents_dir: str = None) -> int:
        """Build the vector knowledge base from documents with optimization"""
        if documents_dir is None:
            documents_dir = self.config.documents_dir
        
        self.logger.info(f"Building knowledge base from {documents_dir}")
        
        try:
            # Process documents
            documents = self.document_processor.process_directory(documents_dir)
            if not documents:
                self.logger.error("No documents processed")
                return 0
            
            # Split documents into chunks with progress tracking
            self.logger.info("Splitting documents into chunks...")
            chunked_docs = []
            
            for doc in tqdm(documents, desc="Chunking documents"):
                chunks = self.text_splitter.split_documents([doc])
                chunked_docs.extend(chunks)
                
                # Memory optimization: process in batches
                if self.config.memory_optimization and len(chunked_docs) > 1000:
                    self._monitor_memory()
            
            self.logger.info(f"Created {len(chunked_docs)} chunks from {len(documents)} documents")
            
            # Create vector store with batch processing
            self.logger.info("Creating vector embeddings...")
            self.vectorstore = FAISS.from_documents(chunked_docs, self.embeddings)
            
            # Save vector store
            self.vectorstore.save_local(self.config.vector_store_path)
            self.logger.info(f"Vector store saved to {self.config.vector_store_path}")
            
            self.logger.info("Knowledge base built successfully")
            
            return len(chunked_docs)
            
        except Exception as e:
            self.logger.error(f"Failed to build knowledge base: {e}")
            return 0
        finally:
            self._monitor_memory()
    
    def load_knowledge_base(self) -> bool:
        """Load existing vector knowledge base with error handling"""
        try:
            if self.embeddings is None:
                self.initialize_models()
            
            # Check if vector store exists
            vector_store_path = Path(self.config.vector_store_path)
            if not vector_store_path.exists():
                self.logger.info("Vector store does not exist")
                return False
            
            self.vectorstore = FAISS.load_local(
                self.config.vector_store_path, 
                self.embeddings,
                allow_dangerous_deserialization=True
            )
            self.logger.info(f"Vector store loaded from {self.config.vector_store_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to load vector store: {e}")
            return False
    
    def setup_rag_chain(self):
        """Setup the RAG chain for question answering"""
        if self.vectorstore is None:
            self.logger.error("Vector store not initialized. Call build_knowledge_base() first.")
            return
        
        # Create retriever
        retriever = self.vectorstore.as_retriever(
            search_type="similarity",
            search_kwargs={'k': self.config.retriever_k}
        )
        
        # Create prompt template for VoWiFi property extraction
        prompt_template = """<|system|>
You are a VoWiFi (Voice over WiFi) expert specializing in property extraction from technical specifications. 
Your task is to extract and identify specific properties, requirements, and behaviors from VoWiFi documentation.

A property is a statement that describes:
1. Requirements (using "shall", "must", "required")
2. Conditional behaviors (starting with "if", "when", "unless")
3. Technical specifications and procedures
4. Security and authentication requirements
5. Message exchange protocols
6. System constraints and limitations

Focus on extracting precise, actionable properties that define how VoWiFi systems should behave.

Context from VoWiFi specifications:
{context}

</s>
<|user|>
{question}
</s>
<|assistant|>"""
        
        prompt = PromptTemplate(
            input_variables=["context", "question"],
            template=prompt_template,
        )
        
        # Create RAG chain
        self.rag_chain = (
            {"context": retriever, "question": RunnablePassthrough()}
            | prompt 
            | self.llm 
            | StrOutputParser()
        )
        
        self.logger.info("RAG chain setup completed")
    
    def extract_properties_from_documents(self, output_file: str = None) -> List[str]:
        """Extract all properties from the knowledge base"""
        if self.vectorstore is None:
            self.logger.error("Vector store not initialized")
            return []
        
        self.logger.info("Extracting properties from all documents...")
        
        # Get all documents from vector store
        all_docs = self.vectorstore.docstore._dict.values()
        all_properties = []
        
        # Use batch processing for large document sets
        docs_list = list(all_docs)
        batch_size = self.config.batch_size
        
        for i in tqdm(range(0, len(docs_list), batch_size), desc="Processing document batches"):
            batch_docs = docs_list[i:i + batch_size]
            batch_properties = []
            
            for doc in batch_docs:
                properties = self.property_extractor.extract_properties_from_text(doc.page_content)
                batch_properties.extend(properties)
            
            # Remove duplicates within batch
            for prop in batch_properties:
                if prop not in all_properties:
                    all_properties.append(prop)
            
            # Memory optimization
            if self.config.memory_optimization and i % (batch_size * 10) == 0:
                self._monitor_memory()
        
        # Save properties
        if output_file is None:
            output_file = os.path.join(self.config.output_dir, "vowifi_properties.txt")
        
        with open(output_file, 'w', encoding='utf-8') as f:
            for i, prop in enumerate(all_properties, 1):
                f.write(f"{i}. {prop}\n")
        
        self.logger.info(f"Extracted {len(all_properties)} properties to {output_file}")
        return all_properties
    
    def query(self, question: str) -> str:
        """Query the RAG system"""
        if self.rag_chain is None:
            self.setup_rag_chain()
        
        try:
            response = self.rag_chain.invoke(question)
            return response.strip()
        except Exception as e:
            self.logger.error(f"Query failed: {e}")
            return f"Error processing query: {e}"
    
    def extract_properties_for_query(self, query: str) -> List[str]:
        """Extract properties relevant to a specific query"""
        # First get relevant context
        retriever = self.vectorstore.as_retriever(
            search_type="similarity",
            search_kwargs={'k': self.config.retriever_k * 2}  # Get more context
        )
        
        relevant_docs = retriever.get_relevant_documents(query)
        combined_context = "\n".join([doc.page_content for doc in relevant_docs])
        
        # Extract properties from relevant context
        properties = self.property_extractor.extract_properties_from_text(combined_context)
        return properties


def main():
    """Main function to extract VoWiFi properties from RFC documents"""
    print("VoWiFi Property Extraction Framework")
    print("=" * 50)
    
    try:
        # Use default configuration
        config = RAGConfig()
        
        # Initialize pipeline
        pipeline = VoWiFiRAGPipeline(config)
        
        # Initialize models
        print("Initializing models...")
        pipeline.initialize_models()
        print("Models initialized successfully")
        
        # Build knowledge base if it doesn't exist
        if not pipeline.load_knowledge_base():
            print("Building knowledge base from RFC documents...")
            chunks_created = pipeline.build_knowledge_base()
            print(f"Knowledge base built with {chunks_created} chunks")
        else:
            print("Loaded existing knowledge base")
        
        # Extract properties
        print("Extracting VoWiFi properties from documents...")
        properties = pipeline.extract_properties_from_documents()
        
        # Simple output
        print(f"Extracted {len(properties)} properties")
        print(f"Properties saved to: {os.path.join(config.output_dir, 'vowifi_properties.txt')}")
        
        return len(properties)
        
    except Exception as e:
        print(f"Error: {e}")
        logging.error(f"Main execution failed: {e}", exc_info=True)
        return 0


if __name__ == "__main__":
    result = main()
    if result > 0:
        print(f"Successfully extracted {result} properties!")
    else:
        print("Extraction failed. Check the logs for details.")